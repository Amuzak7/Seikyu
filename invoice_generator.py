"""
請求書 PDF / Word 生成モジュール。
PDF : fpdf2（純 Python）を使用。WeasyPrint + GTK が利用可能な環境では自動切替。
Word: python-docx を使用。
"""
import os
import base64
from pathlib import Path

import database as db

INVOICES_DIR = Path("invoices")
TEMPLATES_DIR = Path("templates")
ACCENT   = "1E4D7B"          # 紺色（16進）
BLUE_RGB = (30, 77, 123)     # 同・RGB

# Windows 日本語フォント候補（優先順）
_JP_REGULAR_FONTS = ["YuGothR.ttc", "meiryo.ttc", "msgothic.ttc"]
_JP_BOLD_FONTS    = ["YuGothB.ttc", "YuGothM.ttc", "meiryo.ttc", "msgothic.ttc"]
_FONT_DIR = r"C:\Windows\Fonts"


# ─── 共通ユーティリティ ──────────────────────────────────────

def _ensure_dirs() -> None:
    INVOICES_DIR.mkdir(exist_ok=True)


def _get_data(invoice_id: int) -> tuple[dict, object]:
    details = db.get_invoice_details(invoice_id)
    if not details:
        raise ValueError(f"請求書ID {invoice_id} が見つかりません。")
    return details, db.get_company_info()


def _make_filename(customer_name: str, issue_date: str | None, invoice_number: str) -> str:
    invalid = r'\/:*?"<>|'
    safe = "".join(c for c in (customer_name or "unknown") if c not in invalid).strip()
    date_str = (issue_date or "").replace("-", "")
    num_str  = invoice_number.replace("-", "_")
    return f"請求書_{safe}_{date_str}_{num_str}"


def _find_font(candidates: list[str]) -> str:
    for name in candidates:
        path = os.path.join(_FONT_DIR, name)
        if os.path.exists(path):
            return path
    raise RuntimeError(
        f"日本語フォントが見つかりません。{_FONT_DIR} に {candidates} のいずれかを配置してください。"
    )


# ─── PDF 生成 ────────────────────────────────────────────────

def generate_invoice_pdf(invoice_id: int) -> str:
    """請求書 PDF を生成して invoices/ に保存し、ファイルパスを返す。"""
    _ensure_dirs()
    details, company = _get_data(invoice_id)

    existing = details.get("pdf_path")
    if existing and os.path.exists(existing):
        return existing

    fname    = _make_filename(details.get("customer_name", "unknown"),
                              details.get("issue_date"), details.get("invoice_number", ""))
    pdf_path = str(INVOICES_DIR / f"{fname}.pdf")

    # WeasyPrint が使えれば HTML テンプレートで生成、なければ fpdf2 にフォールバック
    if _weasyprint_available():
        _pdf_with_weasyprint(details, company, pdf_path)
    elif _fpdf2_available():
        _pdf_with_fpdf2(details, company, pdf_path)
    else:
        raise RuntimeError(
            "PDF生成ライブラリが見つかりません。\n"
            "以下のコマンドを実行してインストールしてください：\n\n"
            "    pip install fpdf2\n\n"
            "または requirements.txt を使って一括インストール：\n\n"
            "    pip install -r requirements.txt"
        )

    db.update_invoice_paths(invoice_id, pdf_path=pdf_path)
    return pdf_path


def _weasyprint_available() -> bool:
    try:
        import weasyprint  # noqa: F401
        return True
    except Exception:
        return False


def _fpdf2_available() -> bool:
    try:
        from fpdf import FPDF  # noqa: F401
        return True
    except ImportError:
        return False


def _pdf_with_weasyprint(details: dict, company, pdf_path: str) -> None:
    from jinja2 import Environment, FileSystemLoader
    from weasyprint import HTML as WH

    logo_url = _logo_data_url(company.logo_path if company else None)
    env      = Environment(loader=FileSystemLoader(str(TEMPLATES_DIR)))
    html_str = env.get_template("invoice_pdf.html").render(
        invoice=details, company=company, logo_data_url=logo_url,
    )
    WH(string=html_str, base_url=str(Path(".").resolve())).write_pdf(pdf_path)


def _logo_data_url(logo_path: str | None) -> str | None:
    if not logo_path or not os.path.exists(logo_path):
        return None
    suffix = Path(logo_path).suffix.lower()
    mime   = "image/jpeg" if suffix in (".jpg", ".jpeg") else "image/png"
    with open(logo_path, "rb") as f:
        return f"data:{mime};base64,{base64.b64encode(f.read()).decode()}"


def _pdf_with_fpdf2(details: dict, company, pdf_path: str) -> None:
    """fpdf2 で日本語請求書 PDF を生成する（GTK 不要）。"""
    try:
        from fpdf import FPDF, XPos, YPos
    except ImportError as e:
        raise RuntimeError(
            "fpdf2 がインストールされていません。\n"
            "以下のコマンドを実行してください：\n\n"
            "    pip install fpdf2\n\n"
            f"（詳細: {e}）"
        ) from e

    # ── フォント ──────────────────────────────────────────────
    font_r = _find_font(_JP_REGULAR_FONTS)
    font_b = _find_font(_JP_BOLD_FONTS)

    pdf = FPDF(orientation="P", unit="mm", format="A4")
    pdf.set_margins(left=18, top=16, right=18)
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.add_page()
    pdf.add_font("JP",  "",  font_r)
    pdf.add_font("JP",  "B", font_b)

    ML = 18.0; MR = 18.0; PW = 210.0
    UW = PW - ML - MR   # 使用幅 174 mm

    BLUE  = (30, 77, 123)
    LBLUE = (239, 243, 248)
    LGRAY = (248, 249, 250)
    GRAY  = (85,  85,  85)
    MGRAY = (170, 170, 170)
    WHITE = (255, 255, 255)
    BLACK = (26,  26,  26)

    def jp(size: float, bold: bool = False) -> None:
        pdf.set_font("JP", "B" if bold else "", size)

    def tc(*rgb) -> None:  pdf.set_text_color(*rgb)
    def fc(*rgb) -> None:  pdf.set_fill_color(*rgb)
    def dc(*rgb) -> None:  pdf.set_draw_color(*rgb)
    def lw(w: float) -> None: pdf.set_line_width(w)

    NR = XPos.RIGHT;   NT = YPos.TOP
    NL = XPos.LMARGIN; NN = YPos.NEXT

    # ── ヘッダー ──────────────────────────────────────────────
    y0 = pdf.get_y()

    if company and company.logo_path and os.path.exists(company.logo_path):
        pdf.image(company.logo_path, x=ML, y=y0, h=13)
    else:
        jp(13, True); tc(*BLUE)
        pdf.set_xy(ML, y0)
        pdf.cell(UW * 0.5, 9, company.company_name if company else "", new_x=NR, new_y=NT)

    jp(22, True); tc(*BLUE)
    pdf.set_xy(ML, y0)
    pdf.cell(UW, 10, "請　求　書", new_x=NL, new_y=NN, align="R")

    if company and company.registration_number:
        jp(8); tc(*GRAY)
        pdf.set_x(ML)
        pdf.cell(UW, 4, f"適格請求書発行事業者 登録番号：{company.registration_number}",
                 new_x=NL, new_y=NN, align="R")

    # 区切り線
    pdf.ln(2)
    yl = pdf.get_y()
    lw(0.5); dc(*BLUE)
    pdf.line(ML, yl, ML + UW, yl)
    pdf.ln(4)

    # ── 請求書番号・日付 ──────────────────────────────────────
    label_w = 24.0
    for label, val in [
        ("請求書番号", details.get("invoice_number", "")),
        ("発行日",     details.get("issue_date", "")),
        ("支払期限",   details.get("due_date", "")),
    ]:
        pdf.set_x(ML)
        jp(9); tc(*GRAY)
        pdf.cell(label_w, 5.5, label, new_x=NR, new_y=NT)
        jp(9, True); tc(*BLACK)
        pdf.cell(UW - label_w, 5.5, val, new_x=NL, new_y=NN)
    pdf.ln(4)

    # ── 顧客情報 + 自社情報 ───────────────────────────────────
    y_p = pdf.get_y()
    half = UW * 0.52

    # 左：顧客
    cust_name = (details.get("customer_name") or "（顧客名なし）") + "　御中"
    jp(14, True); tc(*BLUE)
    pdf.set_xy(ML, y_p)
    pdf.cell(half, 9, cust_name, new_x=NL, new_y=NN)
    yl2 = pdf.get_y() - 1
    lw(0.3); dc(*BLUE)
    pdf.line(ML, yl2, ML + half, yl2)
    if details.get("customer_address"):
        jp(8.5); tc(*GRAY)
        pdf.set_xy(ML, pdf.get_y() + 1)
        pdf.multi_cell(half - 2, 4.5, details["customer_address"])

    # 右：自社情報
    y_co = y_p
    co_x = ML + half + 3
    co_w = UW - half - 3
    if company:
        for text, bold, sz in [
            (company.company_name,      True,  11),
            (company.address or "",     False,  8.5),
            (f"TEL：{company.phone}" if company.phone else "",                    False, 8.5),
            (f"登録番号：{company.registration_number}" if company.registration_number else "", False, 8),
        ]:
            if not text:
                continue
            jp(sz, bold); tc(*BLACK if bold else GRAY)
            pdf.set_xy(co_x, y_co)
            pdf.cell(co_w, 5.5, text, new_x=NL, new_y=NN, align="R")
            y_co = pdf.get_y()

    pdf.set_y(max(pdf.get_y(), y_p + 24))
    pdf.ln(4)

    # ── ご請求金額ハイライト ──────────────────────────────────
    y_h = pdf.get_y()
    fc(*LBLUE); dc(*LBLUE); lw(0)
    pdf.rect(ML, y_h, UW, 9, "F")
    fc(*BLUE)
    pdf.rect(ML, y_h, 1.5, 9, "F")
    jp(13, True); tc(*BLUE)
    pdf.set_xy(ML, y_h)
    pdf.cell(UW - 2, 9, f"ご請求金額（税込）：¥{int(details['total']):,}", new_x=NL, new_y=NN, align="R")
    pdf.ln(5)

    # ── 明細テーブル ──────────────────────────────────────────
    COL_W = [88, 16, 26, 26, 18]
    HDRS  = [("品名","L"),("数量","C"),("単価（円）","R"),("金額（円）","R"),("税率","C")]
    lw(0); fc(*BLUE); dc(*BLUE)
    for (hdr, al), w in zip(HDRS, COL_W):
        jp(8.5, True); tc(*WHITE)
        pdf.cell(w, 7, hdr, border=0, new_x=NR, new_y=NT, align=al, fill=True)
    pdf.ln(7)

    has_8 = False
    for ri, item in enumerate(details.get("items", [])):
        bg = LGRAY if ri % 2 == 0 else WHITE
        fc(*bg); dc(*MGRAY); lw(0.15)
        if item["tax_rate"] == 8:
            has_8 = True
        vals   = [
            item["item_name"], str(item["quantity"]),
            f"¥{int(item['unit_price']):,}", f"¥{int(item['amount']):,}",
            f"{item['tax_rate']}%" + ("※" if item["tax_rate"] == 8 else ""),
        ]
        aligns = ["L","C","R","R","C"]
        jp(9); tc(*BLACK)
        for v, al, w in zip(vals, aligns, COL_W):
            pdf.cell(w, 6.5, v, border="B", new_x=NR, new_y=NT, align=al, fill=True)
        pdf.ln(6.5)

    if has_8:
        pdf.ln(1)
        jp(8); tc(*GRAY)
        pdf.set_x(ML)
        pdf.cell(UW, 4.5, "※ 軽減税率（8%）対象", new_x=NL, new_y=NN)
    pdf.ln(4)

    # ── 合計欄 ────────────────────────────────────────────────
    SP = UW * 0.50; LBW = 52.0; AMW = UW - SP - LBW
    tot_rows: list[tuple[str, float]] = [("小計（税抜）", details["subtotal"])]
    if details.get("tax_10", 0) > 0:
        tot_rows.append(("消費税（10%）",    details["tax_10"]))
    if details.get("tax_8",  0) > 0:
        tot_rows.append(("消費税（8% 軽減）", details["tax_8"]))

    for label, amount in tot_rows:
        jp(9.5); tc(*BLACK); dc(*MGRAY); lw(0.2)
        pdf.set_x(ML + SP)
        pdf.cell(LBW, 6.5, label, border="B", new_x=NR, new_y=NT, align="L")
        pdf.cell(AMW, 6.5, f"¥{int(amount):,}", border="B", new_x=NL, new_y=NN, align="R")

    lw(0); fc(*BLUE); dc(*BLUE)
    pdf.set_x(ML + SP)
    jp(11, True); tc(*WHITE)
    pdf.cell(LBW, 8, "合計（税込）",           fill=True, new_x=NR, new_y=NT, align="L")
    pdf.cell(AMW, 8, f"¥{int(details['total']):,}", fill=True, new_x=NL, new_y=NN, align="R")
    pdf.ln(6)

    # ── 振込先 ────────────────────────────────────────────────
    if company and company.bank_name:
        y_bk = pdf.get_y()
        fc(*LGRAY); dc(*MGRAY); lw(0.3)
        pdf.rect(ML, y_bk, UW, 8, "DF")

        bank = company.bank_name
        if company.bank_branch:   bank += f"　{company.bank_branch}"
        bank += f"　{company.account_type}　{company.account_number}　{company.account_holder}"

        jp(10, True); tc(*BLUE)
        pdf.set_xy(ML + 2, y_bk)
        pdf.cell(22, 8, "【振込先】", new_x=NR, new_y=NT)
        jp(9); tc(*BLACK)
        pdf.cell(UW - 26, 8, bank, new_x=NL, new_y=NN)
        pdf.ln(5)

    # ── 備考 ──────────────────────────────────────────────────
    if details.get("notes"):
        jp(10, True); tc(*BLACK)
        pdf.set_x(ML)
        pdf.cell(UW, 6, "備考", new_x=NL, new_y=NN)
        jp(9.5); tc(*BLACK); dc(*MGRAY); lw(0.3)
        pdf.set_x(ML)
        pdf.multi_cell(UW, 5.5, details["notes"], border=1)

    pdf.output(pdf_path)


# ─── Word 生成（python-docx） ────────────────────────────────

def generate_invoice_docx(invoice_id: int) -> str:
    """請求書 Word を生成して invoices/ に保存し、ファイルパスを返す。"""
    _ensure_dirs()
    details, company = _get_data(invoice_id)

    existing = details.get("docx_path")
    if existing and os.path.exists(existing):
        return existing

    try:
        from docx import Document
        from docx.shared import Pt, Mm, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError as e:
        raise RuntimeError(f"python-docx が必要です: {e}") from e

    BLUE  = RGBColor(0x1E, 0x4D, 0x7B)
    GRAY  = RGBColor(0x55, 0x55, 0x55)
    WHITE = RGBColor(0xFF, 0xFF, 0xFF)

    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Mm(210); sec.page_height = Mm(297)
    sec.left_margin = sec.right_margin = Mm(20)
    sec.top_margin = Mm(18); sec.bottom_margin = Mm(22)
    doc.styles["Normal"].font.name = "Meiryo"
    doc.styles["Normal"].font.size = Pt(10)

    # ── ヘッダー ──────────────────────────────────────────────
    htbl = doc.add_table(rows=1, cols=2)
    _no_borders(htbl)
    lc = htbl.rows[0].cells[0]
    lp = lc.paragraphs[0]
    if company and company.logo_path and os.path.exists(company.logo_path):
        lp.add_run().add_picture(company.logo_path, width=Inches(1.8))
    else:
        r = lp.add_run(company.company_name if company else "")
        r.font.size = Pt(13); r.font.bold = True; r.font.color.rgb = BLUE

    rc = htbl.rows[0].cells[1]
    rp = rc.paragraphs[0]; rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = rp.add_run("請　求　書")
    r.font.size = Pt(22); r.font.bold = True; r.font.color.rgb = BLUE
    if company and company.registration_number:
        rp2 = rc.add_paragraph(); rp2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r2 = rp2.add_run(f"適格請求書発行事業者 登録番号：{company.registration_number}")
        r2.font.size = Pt(8); r2.font.color.rgb = GRAY

    sep = doc.add_paragraph(); _para_border_bottom(sep, ACCENT)

    # ── 請求書番号・日付 ──────────────────────────────────────
    mtbl = doc.add_table(rows=3, cols=2); _no_borders(mtbl)
    for i, (label, val) in enumerate([
        ("請求書番号", details.get("invoice_number", "")),
        ("発行日",     details.get("issue_date", "")),
        ("支払期限",   details.get("due_date", "")),
    ]):
        cells = mtbl.rows[i].cells
        cells[0].width = Mm(28)
        _run(cells[0].paragraphs[0], label, Pt(9), color=GRAY)
        _run(cells[1].paragraphs[0], val,   Pt(9), bold=True)
    doc.add_paragraph()

    # ── 顧客 + 自社情報 ───────────────────────────────────────
    ptbl = doc.add_table(rows=1, cols=2); _no_borders(ptbl)
    lc2 = ptbl.rows[0].cells[0]
    cname = (details.get("customer_name") or "（顧客名なし）") + "　御中"
    _run(lc2.paragraphs[0], cname, Pt(14), bold=True, color=BLUE)
    if details.get("customer_address"):
        _run(lc2.add_paragraph(), details["customer_address"], Pt(9), color=GRAY)

    rc2 = ptbl.rows[0].cells[1]; first = True
    if company:
        for text, bold, sz in [
            (company.company_name,      True,  Pt(11)),
            (company.address or "",     False, Pt(9)),
            (f"TEL：{company.phone}" if company.phone else "",                     False, Pt(9)),
            (f"登録番号：{company.registration_number}" if company.registration_number else "", False, Pt(8)),
        ]:
            if not text: continue
            p = rc2.paragraphs[0] if first else rc2.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            _run(p, text, sz, bold=bold, color=GRAY if not bold else None)
            first = False
    doc.add_paragraph()

    # ── ご請求金額 ────────────────────────────────────────────
    hp = doc.add_paragraph(); hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _shading_para(hp, "EFF3F8")
    _run(hp, f"ご請求金額（税込）：¥{int(details['total']):,}", Pt(14), bold=True, color=BLUE)
    doc.add_paragraph()

    # ── 明細テーブル ──────────────────────────────────────────
    items = details.get("items", [])
    col_w = [Mm(88), Mm(16), Mm(25), Mm(25), Mm(16)]
    HDRS2 = [("品名","LEFT"),("数量","CENTER"),("単価（円）","RIGHT"),("金額（円）","RIGHT"),("税率","CENTER")]

    itbl = doc.add_table(rows=1 + len(items), cols=5)
    itbl.style = "Table Grid"
    _tbl_borders(itbl, ACCENT)

    for ci, ((hdr, al), w) in enumerate(zip(HDRS2, col_w)):
        cell = itbl.rows[0].cells[ci]; cell.width = w
        _cell_bg(cell, ACCENT)
        p = cell.paragraphs[0]; p.alignment = getattr(WD_ALIGN_PARAGRAPH, al)
        _run(p, hdr, Pt(9), bold=True, color=WHITE)

    has_8 = False
    for ri, item in enumerate(items):
        bg = "F8F9FA" if ri % 2 == 0 else "FFFFFF"
        row = itbl.rows[ri + 1]
        if item["tax_rate"] == 8: has_8 = True
        vals   = [item["item_name"], str(item["quantity"]),
                  f"¥{int(item['unit_price']):,}", f"¥{int(item['amount']):,}",
                  f"{item['tax_rate']}%" + ("※" if item["tax_rate"] == 8 else "")]
        aligns = ["LEFT","CENTER","RIGHT","RIGHT","CENTER"]
        for ci, (v, al, w) in enumerate(zip(vals, aligns, col_w)):
            cell = row.cells[ci]; cell.width = w
            _cell_bg(cell, bg)
            p = cell.paragraphs[0]; p.alignment = getattr(WD_ALIGN_PARAGRAPH, al)
            _run(p, v, Pt(9))

    if has_8:
        np_ = doc.add_paragraph("※ 軽減税率（8%）対象")
        np_.runs[0].font.size = Pt(8); np_.runs[0].font.color.rgb = GRAY
    doc.add_paragraph()

    # ── 合計欄 ────────────────────────────────────────────────
    tot_rows: list[tuple[str, float]] = [("小計（税抜）", details["subtotal"])]
    if details.get("tax_10", 0) > 0: tot_rows.append(("消費税（10%）",    details["tax_10"]))
    if details.get("tax_8",  0) > 0: tot_rows.append(("消費税（8% 軽減）", details["tax_8"]))

    SP_MM = 88; LB_MM = 46; AM_MM = 36
    ttbl = doc.add_table(rows=len(tot_rows) + 1, cols=3); _no_borders(ttbl)
    for ri, (label, amount) in enumerate(tot_rows):
        cells = ttbl.rows[ri].cells
        cells[0].width = Mm(SP_MM); cells[1].width = Mm(LB_MM); cells[2].width = Mm(AM_MM)
        _run(cells[1].paragraphs[0], label, Pt(9.5))
        p = cells[2].paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _run(p, f"¥{int(amount):,}", Pt(9.5))

    gc = ttbl.rows[-1].cells
    gc[0].width = Mm(SP_MM); gc[1].width = Mm(LB_MM); gc[2].width = Mm(AM_MM)
    _cell_bg(gc[1], ACCENT); _cell_bg(gc[2], ACCENT)
    _run(gc[1].paragraphs[0], "合計（税込）", Pt(11), bold=True, color=WHITE)
    p2 = gc[2].paragraphs[0]; p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _run(p2, f"¥{int(details['total']):,}", Pt(11), bold=True, color=WHITE)
    doc.add_paragraph()

    # ── 振込先 ────────────────────────────────────────────────
    if company and company.bank_name:
        bp = doc.add_paragraph(); _shading_para(bp, "FAFAFA")
        _run(bp, "【振込先】　", Pt(10), bold=True, color=BLUE)
        bank = company.bank_name
        if company.bank_branch: bank += f"　{company.bank_branch}"
        bank += f"　{company.account_type}　{company.account_number}　{company.account_holder}"
        _run(bp, bank, Pt(9.5))
        doc.add_paragraph()

    # ── 備考 ──────────────────────────────────────────────────
    if details.get("notes"):
        np_ = doc.add_paragraph(); _run(np_, "備考", Pt(10), bold=True)
        for line in details["notes"].split("\n"):
            p = doc.add_paragraph(line)
            if p.runs: p.runs[0].font.size = Pt(9.5)

    fname    = _make_filename(details.get("customer_name", "unknown"),
                              details.get("issue_date"), details.get("invoice_number", ""))
    docx_path = str(INVOICES_DIR / f"{fname}.docx")
    doc.save(docx_path)
    db.update_invoice_paths(invoice_id, docx_path=docx_path)
    return docx_path


# ─── python-docx XML ヘルパー ────────────────────────────────

def _run(para, text: str, size, bold: bool = False, color=None):
    r = para.add_run(text)
    r.font.size = size; r.font.bold = bold
    if color is not None: r.font.color.rgb = color
    return r


def _no_borders(table) -> None:
    from docx.oxml.ns import qn; from docx.oxml import OxmlElement
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None: tblPr = OxmlElement("w:tblPr"); tbl.insert(0, tblPr)
    bd = OxmlElement("w:tblBorders")
    for name in ["top","left","bottom","right","insideH","insideV"]:
        b = OxmlElement(f"w:{name}"); b.set(qn("w:val"), "none"); bd.append(b)
    tblPr.append(bd)


def _tbl_borders(table, accent: str) -> None:
    from docx.oxml.ns import qn; from docx.oxml import OxmlElement
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None: tblPr = OxmlElement("w:tblPr"); tbl.insert(0, tblPr)
    bd = OxmlElement("w:tblBorders")
    for name in ["top","left","bottom","right"]:
        b = OxmlElement(f"w:{name}"); b.set(qn("w:val"),"single"); b.set(qn("w:sz"),"6"); b.set(qn("w:color"),accent); bd.append(b)
    for name in ["insideH","insideV"]:
        b = OxmlElement(f"w:{name}"); b.set(qn("w:val"),"single"); b.set(qn("w:sz"),"4"); b.set(qn("w:color"),"CCCCCC"); bd.append(b)
    tblPr.append(bd)


def _cell_bg(cell, fill: str) -> None:
    from docx.oxml.ns import qn; from docx.oxml import OxmlElement
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd"); shd.set(qn("w:val"),"clear"); shd.set(qn("w:color"),"auto"); shd.set(qn("w:fill"),fill)
    tcPr.append(shd)


def _shading_para(para, fill: str) -> None:
    from docx.oxml.ns import qn; from docx.oxml import OxmlElement
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd"); shd.set(qn("w:val"),"clear"); shd.set(qn("w:color"),"auto"); shd.set(qn("w:fill"),fill)
    pPr.append(shd)


def _para_border_bottom(para, color: str) -> None:
    from docx.oxml.ns import qn; from docx.oxml import OxmlElement
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom"); bot.set(qn("w:val"),"single"); bot.set(qn("w:sz"),"12"); bot.set(qn("w:color"),color)
    pBdr.append(bot); pPr.append(pBdr)
